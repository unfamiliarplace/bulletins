from typing import List, Any, Optional, Union, Generator
from copy import deepcopy
import re
import os

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.table import Table, _Row, _Cell
from docx.text.paragraph import Paragraph


class DocumentTools:

    @staticmethod
    def copy_table(table: Table) -> Table:
        new = deepcopy(table)
        new._tbl = deepcopy(table._tbl)
        return new

    @staticmethod
    def copy_paragraph(paragraph: Paragraph) -> Paragraph:
        new = deepcopy(paragraph)
        new._p = deepcopy(paragraph._p)
        return new

    @staticmethod
    def _delete_element(proxy: Union[Paragraph, Table]) -> None:
        el = proxy._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    @staticmethod
    def delete_paragraph(paragraph: Paragraph) -> None:
        DocumentTools._delete_element(paragraph)
        paragraph._p = paragraph._element = None

    @staticmethod
    def delete_table(table: Table) -> None:
        DocumentTools._delete_element(table)
        table._tbl = table._element = None

    @staticmethod
    def delete_row(table: Table, row_i: int) -> None:
            """
            https://stackoverflow.com/a/71336349
            """
            row = table.rows[row_i]
            table_element = table._tbl
            table_element.remove(row._tr)

    @staticmethod
    def paragraph_text(p: Paragraph) -> str:
        return ''.join(run.text for run in p.runs)

    @staticmethod
    def shade_row(r: _Row, colour: str) -> None:
        for cell in r.cells:
            cell._tc.get_or_add_tcPr().append(DocumentTools.get_shade(colour))

    @staticmethod
    def get_shade(colour: str) -> None:
        w = 'w'
        return parse_xml(f'<w:shd {nsdecls(w)} w:fill="{colour}"/>')

    @staticmethod
    def row_cell(row: _Row, i_cell: int) -> str:
        return '\n'.join(p.text for p in row.cells[i_cell].paragraphs)

    @staticmethod
    def row_para(row: _Row, i_cell: int, i_para: int) -> str:
        return row.cells[i_cell].paragraphs[i_para].text

    @staticmethod
    def cell_para(cell: _Cell, i_para: int) -> str:
        return cell.paragraphs[i_para].text

    @staticmethod
    def add_row(t: Table, cells: List[Any], colour: str = '') -> _Row:
        row = t.add_row()
        for i in range(len(cells)):
            row.cells[i].paragraphs[0].text = str(cells[i])

        if colour:
            DocumentTools.shade_row(row, colour)

        return row

    @staticmethod
    def find_paragraph(d: Document, pattern: str, in_text: bool = True,
                       in_table: bool = True) -> Optional[Paragraph]:
        gen = DocumentTools._find_paragraphs(d, pattern, in_text, in_table)
        for p in gen:
            return p  # Yup, intended as a break

    @staticmethod
    def find_paragraphs(d: Document, pattern: str, in_text: bool = True,
                        in_table: bool = True) -> List[Paragraph]:

        gen = DocumentTools._find_paragraphs(d, pattern, in_text, in_table)
        return list(gen)

    @staticmethod
    def _find_paragraphs(d: Document, pattern: str, in_text: bool = True,
                         in_table: bool = True) -> Generator[Paragraph, None, None]:

        if in_text:
            for p in d.paragraphs:
                matches = re.findall(pattern, p.text, re.MULTILINE)
                if matches:
                    yield p

        if in_table:
            for t in d.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            matches = re.findall(pattern, p.text, re.MULTILINE)
                            if matches:
                                yield p

    @staticmethod
    def sub(d: Document, key: str, value: str, stop: bool = True, in_text: bool=True, in_table: bool=True) -> None:
        for p in DocumentTools._find_paragraphs(d, key, in_text=in_text, in_table=in_table):
            p.text = p.text.replace(key, str(value))
            if stop:
                return

    @staticmethod
    def sub_styles(d: Document) -> None:
        re_style = r'(__style: (.*?)__)'

        for p in DocumentTools._find_paragraphs(d, re_style):
            matches = re.findall(re_style, p.text, re.MULTILINE)
            p.text = p.text.replace(matches[0][0], '')
            p.style = d.styles[matches[0][1]]

    @staticmethod
    def combine_documents(documents: List[Document]) -> Document:
        """https://stackoverflow.com/a/40490622/5228348"""
        merged_doc = Document()

        for index, doc in enumerate(documents):
            temp_path = '.temp_copy_docx'
            doc.save(temp_path)
            doc2 = Document(temp_path)

            # Don't add a page break if you've reached the last file
            if index < len(documents) - 1:
                doc2.add_page_break()

            for element in doc2.element.body:
                merged_doc.element.body.append(element)

            os.remove(temp_path)

        return merged_doc

    @staticmethod
    def get_text_width(document: Document) -> int:
        """
        Returns the text width in mm.
        """
        section = document.sections[0]
        return (section.page_width - section.left_margin - section.right_margin) / 36000
