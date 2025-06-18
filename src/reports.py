from pathlib import Path
from typing import List, Dict, Tuple
from docx import Document
from doctools import DocumentTools

PATH_SRC = Path('src')
PATH_TEMPLATES = PATH_SRC / 'templates'
PATH_TEMPLATE_STUDENT = PATH_TEMPLATES / 'student.docx'

class Reports:

    @staticmethod
    def list_files(path: str) -> List[str]:
        return list(filter(lambda p: '~' not in p.stem, path.glob('*.docx')))

    @staticmethod
    def parse_files(paths: List[Path]) -> Dict[str, Tuple[str, Dict[str, str]]]:
        # Returns dict of {date: (question, {student: answer})}

        dates = {}
        for path in paths:            
            ds = path.stem
            question, answers = Reports.parse_file(Document(path))
            dates[ds] = (question, answers)

        return dates

    @staticmethod
    def parse_file(doc: Document) -> Tuple[str, Dict[str, str]]:
        question = ''
        answers = {}

        state = 0
        for para in doc.paragraphs:
            text = DocumentTools.paragraph_text(para)

            if state == 0:
                question = text
                state = 1

            elif state == 1:
                if text.strip():
                    question += '\n' + text
                else:
                    state = 2

            elif state == 2:
                if text:
                    student = text[:text.find(':')]
                    answer = text[text.find(':')+1:].strip()
                    answers[student] = answer

        return question, answers

    @staticmethod
    def map_aliases(path_dir: Path) -> dict[str, str]:
        d = {}
        path_names = path_dir / '_names.txt'
        if not path_names.exists():
            return d
        
        with open(path_names, 'r', encoding='utf-8') as f:
            for line in f.readlines():
                line = line.strip()
                if line and not line.startswith(';'):
                    alias, real = line.split('::')
                    d[alias.lower()] = real

        return d

    @staticmethod
    def collate_students(dses: Dict[str, Tuple[str, Dict[str, str]]], aliases: dict[str, str]) -> Dict[str, List[Tuple[str, str, str]]]:
        students = {}

        for ds in sorted(dses):
            question, answers = dses[ds]

            for (student, answer) in answers.items():
                student = student.strip()

                student = aliases.get(student.lower(), student)

                if student == '-' or student.strip() == '':
                    continue

                if answer.strip() == '-':
                    continue
                
                if student not in students:
                    students[student] = []

                students[student].append((ds, question, answer))

        return students

    @staticmethod
    def make_documents(students: Dict[str, List[Tuple[str, str, str]]]) -> Dict[str, Document]:
        docs = {}

        for student in students:
            doc = Reports.make_document(student, students[student])
            docs[student] = doc

        return docs

    @staticmethod
    def make_document(student: str, dses: List[Tuple[str, str, str]]) -> Document:
        d = Document(PATH_TEMPLATE_STUDENT)

        DocumentTools.sub(d, '__student__', student)

        for (ds, question, answer) in dses:
            # DocumentTools.add_row(table, [ds, question, answer])
            q = d.add_paragraph(f'({ds}): {question}')
            a = d.add_paragraph(answer)
            b = d.add_paragraph('')

            q.style = 'Question'
            a.style = 'Answer'
            b.style = 'Space'

        return d

    @staticmethod
    def save_documents(docs: Dict[str, Document], path_output_dir: Path) -> None:
        if not path_output_dir.exists():
            path_output_dir.mkdir(parents=True)

        for ds in docs:
            path = path_output_dir / (f'{ds}.docx')
            docs[ds].save(path)
