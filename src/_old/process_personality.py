from __future__ import annotations

import os
import re
import subprocess
from copy import deepcopy
from tkinter import Tk, messagebox
from tkinter.filedialog import askdirectory
from typing import List, Dict, Any, Optional, Union, Generator, Tuple, Set, TextIO

# package name is a bit different. pip install python-docx
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.table import Table, _Row, _Cell
from docx.text.paragraph import Paragraph

INPUT_ROOT = 'input'
OUTPUT_ROOT = 'output'
STUDENT_TEMPLATE = 'templates/student.docx'

NAMES = {
    # example
    #'Josh': 'Joshua',

    # 2021-22 S1 ENG4U
##    'Kennis': 'Ken',
##    'Denis de Villeneuve': 'Dennis',
##    'Luc de Boer': 'Luke D',
##    'Luke DeBoer': 'Luke D',
##    'Luke': 'Luke B',
##    'Josh C': 'Josh',
##    'Joshua': 'Josh',
##    'Josh M': 'Josh',
##    'Wins-a-ton': 'Winston',
##    'Tori': 'Victoria',

    # 2021-22 S1 FSF1D
##    'Allleexxxx': 'Alex',
##    'Daniel': 'Daniel K',
##    'Delilah': 'Daliah',
##    'Ericccc': 'Eric',
##    'Jacob': 'Jacob Z',
##    'Jacob D': 'Jake D',
##    'Jake': 'Jake D',
##    'Neveah': 'Nevaeh',
##    'Sammy S': 'Samy S',
##    'Sammy': 'Samy S',
##    'Samuel': 'Samuel C',
##    'Samy': 'Samy S'

    # 2021-22 S1 ICS3U
    'Balex': 'Alex B',
    'Daniél': 'Daniel',
    'Fran "Seen"': 'Francine',
    'Frâncïné': 'Francine',
    'Francyan': 'Francine',
    'Katarina': 'Kathy',
    'Mark “Kania Believe It” Kania': 'Mark',
    'Marcus Aurelius': 'Mark',
    'Pablo': 'Paul',
    'Matt': 'Matthew',
    'Markham': 'Mark',
    'Roána': 'Rona',
    'Ryu': 'Luke',
    'Thomász': 'Thomas',
    'Tomász': 'Thomas',
    'Gianmarco': 'Mark',
}

class DocumentTools:

    @staticmethod
    def copy_table(table: Table) -> Table:
        new = deepcopy(table)
        new._tbl = deepcopy(table._tbl)
        return new

    @staticmethod
    def copy_paragraph(paragraph: Paragraph) -> Paragraph:
        new = deepcopy(paragraph)
        new._p = deepcopy(paragraph._p)
        return new

    @staticmethod
    def _delete_element(proxy: Union[Paragraph, Table]) -> None:
        el = proxy._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    @staticmethod
    def delete_paragraph(paragraph: Paragraph) -> None:
        DocumentTools._delete_element(paragraph)
        paragraph._p = paragraph._element = None

    @staticmethod
    def delete_table(table: Table) -> None:
        DocumentTools._delete_element(table)
        table._tbl = table._element = None

    @staticmethod
    def paragraph_text(p: Paragraph) -> str:
        return ''.join(run.text for run in p.runs)

    @staticmethod
    def shade_row(r: _Row, colour: str) -> None:
        for cell in r.cells:
            cell._tc.get_or_add_tcPr().append(DocumentTools.get_shade(colour))

    @staticmethod
    def get_shade(colour: str) -> None:
        w = 'w'
        return parse_xml(f'<w:shd {nsdecls(w)} w:fill="{colour}"/>')

    @staticmethod
    def row_cell(row: _Row, i_cell: int) -> str:
        return '\n'.join(p.text for p in row.cells[i_cell].paragraphs)

    @staticmethod
    def row_para(row: _Row, i_cell: int, i_para: int) -> str:
        return row.cells[i_cell].paragraphs[i_para].text

    @staticmethod
    def cell_para(cell: _Cell, i_para: int) -> str:
        return cell.paragraphs[i_para].text

    @staticmethod
    def add_row(t: Table, cells: List[Any], colour: str = '') -> _Row:
        row = t.add_row()
        for i in range(len(cells)):
            row.cells[i].paragraphs[0].text = str(cells[i])

        if colour:
            DocumentTools.shade_row(row, colour)

        return row

    @staticmethod
    def find_paragraph(d: Document, pattern: str, in_text: bool = True,
                       in_table: bool = True) -> Optional[Paragraph]:
        gen = DocumentTools._find_paragraphs(d, pattern, in_text, in_table)
        for p in gen:
            return p  # Yup, intended as a break

    @staticmethod
    def find_paragraphs(d: Document, pattern: str, in_text: bool = True,
                        in_table: bool = True) -> List[Paragraph]:

        gen = DocumentTools._find_paragraphs(d, pattern, in_text, in_table)
        return list(gen)

    @staticmethod
    def _find_paragraphs(d: Document, pattern: str, in_text: bool = True,
                         in_table: bool = True) -> Generator[Paragraph]:

        if in_text:
            for p in d.paragraphs:
                matches = re.findall(pattern, p.text, re.MULTILINE)
                if matches:
                    yield p

        if in_table:
            for t in d.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            matches = re.findall(pattern, p.text, re.MULTILINE)
                            if matches:
                                yield p

    @staticmethod
    def sub(d: Document, key: str, value: str, stop: bool = True) -> None:
        for p in DocumentTools._find_paragraphs(d, key):
            p.text = p.text.replace(key, value)
            if stop:
                return

    @staticmethod
    def sub_styles(d: Document) -> None:
        re_style = r'(__style: (.*?)__)'

        for p in DocumentTools._find_paragraphs(d, re_style):
            matches = re.findall(re_style, p.text, re.MULTILINE)
            p.text = p.text.replace(matches[0][0], '')
            p.style = d.styles[matches[0][1]]

    @staticmethod
    def combine_documents(documents: List[Document]) -> Document:
        """https://stackoverflow.com/a/40490622/5228348"""
        merged_doc = Document()

        for index, doc in enumerate(documents):
            temp_path = '.temp_copy_docx'
            doc.save(temp_path)
            doc2 = Document(temp_path)

            # Don't add a page break if you've reached the last file
            if index < len(documents) - 1:
                doc2.add_page_break()

            for element in doc2.element.body:
                merged_doc.element.body.append(element)

            os.remove(temp_path)

        return merged_doc


class Reports:

    @staticmethod
    def list_files(path: str) -> List[str]:
        return list(os.path.join(path, p) for p in list(os.walk(path))[0][2])

    @staticmethod
    def parse_files(filenames: List[str]) -> Dict[str, Tuple[str, Dict[str, str]]]:
        # Returns dict of {date: (question, {student: answer})}

        dates = {}
        for filename in filenames:
            if '\~' in filename:
                continue

            
            ds = filename[filename.rfind('\\')+1:][:-5]
            question, answers = Reports.parse_file(Document(filename))
            dates[ds] = (question, answers)

        return dates

    @staticmethod
    def parse_file(doc: Document) -> Tuple[str, Dict[str, str]]:
        question = ''
        answers = {}

        state = 0
        for para in doc.paragraphs:
            text = DocumentTools.paragraph_text(para)

            if state == 0:
                question = text
                state = 1

            elif state == 1:
                if text.strip():
                    question += '\n' + text
                else:
                    state = 2

            elif state == 2:
                if text:
                    student = text[:text.find(':')]
                    answer = text[text.find(':')+1:].strip()
                    answers[student] = answer

        return question, answers

    @staticmethod
    def collate_students(dses: Dict[str, Tuple[str, Dict[str, str]]]) -> Dict[str, List[Tuple[str, str, str]]]:
        students = {}

        for ds in dses:
            question, answers = dses[ds]

            for (student, answer) in answers.items():
                student = student.strip()
                
                student = NAMES[student] if student in NAMES else student
                if student not in students:
                    students[student] = []

                students[student].append((ds, question, answer))

        return students

    @staticmethod
    def make_documents(students: Dict[str, List[Tuple[str, str, str]]]) -> Dict[str, Document]:
        docs = {}

        for student in students:
            doc = Reports.make_document(student, students[student])
            docs[student] = doc

        return docs

    @staticmethod
    def make_document(student: str, dses: List[Tuple[str, str, str]]) -> Document:
        d = Document(STUDENT_TEMPLATE)

        DocumentTools.sub(d, '__student__', student)

        table = d.tables[0]
        for ds in dses:
            DocumentTools.add_row(table, [ds[0], ds[1], ds[2]])

        return d

    @staticmethod
    def save_documents(docs: Dict[str, Document]) -> None:
        for ds in docs:
            path = os.path.join(OUTPUT_ROOT, ds) + '.docx'
            docs[ds].save(path)


def prompt_directory() -> str:
    return askdirectory(title='Input directory', initialdir=os.getcwd())


def show_result(path: str) -> None:
    msg = f'Reports placed in\n{path}/\n\nOpen folder now?'
    if messagebox.askyesno('Reports generated', msg):
        open_path = path.replace('/', '\\')
        subprocess.Popen(f'explorer "{open_path}"')


def generate_reports() -> None:
    input_directory = prompt_directory()
    filenames = Reports.list_files(input_directory)    
    dses = Reports.parse_files(filenames)
    students = Reports.collate_students(dses)
    documents = Reports.make_documents(students)
    Reports.save_documents(documents)
    show_result(OUTPUT_ROOT)


if __name__ == '__main__':
    root = Tk()
    root.title('TDChristian Personality Questionnaire')
    try:
        generate_reports()
    except Exception as e:
        messagebox.showerror(title='Error', message=f'Could not process.\n{repr(e)}')
    finally:
        root.destroy()
